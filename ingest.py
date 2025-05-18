import os
from docx import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
#from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LangchainDocument
import config

def load_docx_file(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [LangchainDocument(page_content=text)]

def load_all_documents(folder_path):
    documents = []
    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            loader = PyPDFLoader(path)
            docs = loader.load()
        elif filename.endswith(".txt") or filename.endswith(".md"):
            loader = TextLoader(path)
            docs = loader.load()
        elif filename.endswith(".docx"):
            docs = load_docx_file(path)
        else:
            print(f"❌ Skipping unsupported file: {filename}")
            continue
        documents.extend(docs)
    return documents

def ingest_documents():
    print("📥 Loading documents...")
    documents = load_all_documents("data")

    print("🔪 Splitting into chunks...")
    splitter = RecursiveCharacterTextSplitter(chunk_size=config.CHUNK_SIZE, chunk_overlap=config.CHUNK_OVERLAP)
    chunks = splitter.split_documents(documents)

    print(f"🔤 Generating embeddings for {len(chunks)} chunks...")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    print("📦 Building FAISS index...")
    vectorstore = FAISS.from_documents(chunks, embeddings)

    print(f"💾 Saving index to: {config.VECTORSTORE_DIR}")
    vectorstore.save_local(config.VECTORSTORE_DIR)

if __name__ == "__main__":
    ingest_documents()
